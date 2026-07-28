import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 27, 75)  # #1E1B4B
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(79, 70, 229)  # #4F46E5
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 65, 85)  # #334155
    return p

def add_body_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(30, 41, 59)  # #1E293B
    return p

def add_code_snippet(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    
    # Background border
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(15, 23, 42)
    return p

def add_figure_image(doc, img_path, fig_num, caption_title, description):
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(6.0))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(6)
        run_fig = p_cap.add_run(f"Figure {fig_num}: ")
        run_fig.bold = True
        run_fig.font.size = Pt(10)
        run_fig.font.name = 'Arial'
        run_title = p_cap.add_run(caption_title)
        run_title.italic = True
        run_title.font.size = Pt(10)
        run_title.font.name = 'Arial'

        add_body_p(doc, description)

def create_table_custom(doc, headers, data, tbl_num, caption_title):
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(10)
    p_cap.paragraph_format.space_after = Pt(4)
    run_tbl = p_cap.add_run(f"Table {tbl_num}: ")
    run_tbl.bold = True
    run_tbl.font.size = Pt(10)
    run_tbl.font.name = 'Arial'
    run_title = p_cap.add_run(caption_title)
    run_title.italic = True
    run_title.font.size = Pt(10)
    run_title.font.name = 'Arial'

    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], '1E293B')
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9.5)
                r.font.name = 'Arial'

    # Format Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = 'F8FAFC' if r_idx % 2 == 1 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            for p in row_cells[c_idx].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Calibri'
                    r.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

print("Helper functions compiled.")
